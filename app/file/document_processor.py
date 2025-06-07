from typing import List
from pathlib import Path
from io import BytesIO
import logging

import docx
import pdfplumber
import PyPDF2

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing various document types and extracting text."""

    def __init__(self, chunk_size: int = 3000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF bytes with robust error handling using multiple methods."""
        logger.info(f"Processing PDF of size: {len(file_content)} bytes")

        # Method 1: Try pdfplumber (most robust)
        try:
            logger.info("Attempting PDF extraction with pdfplumber...")
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                text = ""
                logger.info(f"PDF has {len(pdf.pages)} pages")

                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                            logger.debug(
                                f"Extracted {len(page_text)} characters from page {page_num + 1}"
                            )
                    except Exception as page_error:
                        logger.warning(
                            f"pdfplumber: Error on page {page_num + 1}: {str(page_error)}"
                        )
                        continue

                if text.strip():
                    logger.info(
                        f"pdfplumber: Successfully extracted {len(text)} characters"
                    )
                    return text.strip()
                else:
                    logger.warning("pdfplumber: No text extracted, trying PyPDF2...")

        except Exception as e:
            logger.warning(f"pdfplumber failed: {str(e)}, trying PyPDF2...")

        # Method 2: Fallback to PyPDF2
        try:
            logger.info("Attempting PDF extraction with PyPDF2...")
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ""

            logger.info(f"PDF has {len(pdf_reader.pages)} pages")

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    logger.info(f"PyPDF2: Page {page_num + 1} text: {page_text}")
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                        logger.debug(
                            f"Extracted {len(page_text)} characters from page {page_num + 1}"
                        )
                except Exception as page_error:
                    logger.warning(
                        f"PyPDF2: Error on page {page_num + 1}: {str(page_error)}"
                    )
                    continue

            if text.strip():
                logger.info(f"PyPDF2 Successfully extracted {len(text)} characters")
                return text.strip()
            else:
                raise ValueError(
                    "No readable text found in PDF. The PDF might be image-based, encrypted, or corrupted."
                )

        except Exception as e:
            logger.error(f"PyPDF2 also failed: {str(e)}")
            raise ValueError(
                f"PDF text extraction failed with both methods. The PDF might be image-based, encrypted, password-protected, or corrupted. Error: {str(e)}"
            )

    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX bytes."""
        try:
            logger.info(f"Processing DOCX of size: {len(file_content)} bytes")

            doc = docx.Document(BytesIO(file_content))
            text = ""

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"

            extracted_text = text.strip()

            if not extracted_text:
                raise ValueError("No text content found in DOCX document")

            logger.info(
                f"Successfully extracted {len(extracted_text)} characters from DOCX"
            )
            return extracted_text

        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}")
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")

    def _extract_text_from_txt(self, file_content: bytes) -> str:
        """Extract text from plain text bytes."""
        try:
            logger.info(f"Processing text file of size: {len(file_content)} bytes")

            # Try UTF-8 first
            try:
                text = file_content.decode("utf-8").strip()
            except UnicodeDecodeError:
                # Fall back to latin-1
                try:
                    text = file_content.decode("latin-1").strip()
                except UnicodeDecodeError:
                    # Last resort: UTF-8 with error handling
                    text = file_content.decode("utf-8", errors="replace").strip()

            if not text:
                raise ValueError("Text file appears to be empty")

            logger.info(f"Successfully extracted {len(text)} characters from text file")
            return text

        except Exception as e:
            logger.error(f"Text file processing error: {str(e)}")
            raise ValueError(f"Error extracting text from TXT: {str(e)}")

    def _extract_text_from_file(self, file_content: bytes, file_extension: str) -> str:
        """Extract text from file based on extension."""
        file_extension = file_extension.lower().lstrip(".")

        logger.info(f"Processing file with extension: {file_extension}")

        if file_extension == "pdf":
            return self._extract_text_from_pdf(file_content)
        elif file_extension in ["docx", "doc"]:
            return self._extract_text_from_docx(file_content)
        elif file_extension in ["txt", "md", "py", "js", "html", "css", "json", "xml"]:
            return self._extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        if not text.strip():
            logger.warning("Empty text provided for chunking")
            return []

        words = text.split()
        logger.info(
            f"Chunking text with {len(words)} words (chunk_size={self.chunk_size}, overlap={self.chunk_overlap})"
        )

        if len(words) <= self.chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(words):
            end = min(start + self.chunk_size, len(words))
            chunk_words = words[start:end]
            chunk_text = " ".join(chunk_words)
            chunks.append(chunk_text)

            # Move start position with overlap
            start = end - self.chunk_overlap
            if start >= len(words):
                break

        logger.info(f"Created {len(chunks)} text chunks")
        return chunks

    def process_file(self, file_content: bytes, file_name: str) -> List[str]:
        """Process a file and return text chunks."""
        logger.info(f"Processing file: {file_name}")

        file_extension = Path(file_name).suffix

        # Extract text
        text = self._extract_text_from_file(file_content, file_extension)

        # Validate extracted text
        if not text or len(text.strip()) < 10:
            raise ValueError(
                f"Insufficient text content extracted from {file_name}. Got {len(text)} characters."
            )

        # Split into chunks
        chunks = self._chunk_text(text)

        if not chunks:
            raise ValueError(f"No text chunks could be created from {file_name}")

        logger.info(f"Processed {file_name}: {len(chunks)} chunks created")
        return chunks

    def get_content_type_from_extension(self, extension: str) -> str:
        """Map file extension to content type"""

        content_types = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".txt": "text/plain",
            ".md": "text/markdown",
            ".py": "text/x-python",
            ".js": "text/javascript",
            ".html": "text/html",
            ".css": "text/css",
            ".json": "application/json",
            ".xml": "application/xml",
        }
        return content_types.get(extension, "application/octet-stream")


document_processor = DocumentProcessor()
